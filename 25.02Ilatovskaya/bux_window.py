import sys
from PyQt6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QLabel, QLineEdit, QPushButton,
    QMessageBox, QCheckBox, QFormLayout, QDialog, QDialogButtonBox, QComboBox,
    QDateTimeEdit, QFileDialog, QTableWidget, QTableWidgetItem
)
from PyQt6.QtCore import QTimer, Qt
import pymysql
from docx import Document
from docx.shared import Pt

class AccountantWindow(QWidget):
    def __init__(self, user_data):
        super().__init__()
        self.user_data = user_data
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Бухгалтер')
        self.setGeometry(100, 100, 400, 300)

        layout = QVBoxLayout()

        self.user_label = QLabel(f'Добро пожаловать, {self.user_data[2]} {self.user_data[1]}')
        layout.addWidget(self.user_label)

        self.role_label = QLabel(f'Роль: {self.get_user_role(self.user_data[6])}')
        layout.addWidget(self.role_label)

        self.generate_financial_report_button = QPushButton('Сформировать финансовый отчет')
        self.generate_financial_report_button.clicked.connect(self.generate_financial_report)
        layout.addWidget(self.generate_financial_report_button)

        self.logout_button = QPushButton('Выйти')
        self.logout_button.clicked.connect(self.logout)
        layout.addWidget(self.logout_button)

        self.setLayout(layout)

    def get_user_role(self, role_id):
        try:
            connection = connect_to_database()
            cursor = connection.cursor()
            query = "SELECT name FROM post WHERE id = %s"
            cursor.execute(query, (role_id,))
            result = cursor.fetchone()
            return result[0] if result else 'Неизвестно'
        except pymysql.Error as e:
            QMessageBox.critical(self, 'Ошибка базы данных', f'Ошибка при получении роли: {e}')
            return 'Неизвестно'
        finally:
            cursor.close()
            connection.close()

    def generate_financial_report(self):
        dialog = FinancialReportDialog()
        if dialog.exec():
            data = dialog.get_data()
            file_path = self.create_financial_report(data)
            if file_path:
                QMessageBox.information(self, 'Отчет', f'Финансовый отчет успешно создан и сохранен по пути: {file_path}')

    def create_financial_report(self, data):
        try:
            connection = connect_to_database()
            cursor = connection.cursor()
            query = """
            SELECT ic.name, SUM(i.summa)
            FROM invoices_issued i
            JOIN insurance_company ic ON i.id_company = ic.id
            WHERE i.start_period >= %s AND i.end_period <= %s AND i.id_company = %s
            GROUP BY ic.name
            """
            cursor.execute(query, (data['start_period'], data['end_period'], data['company_id']))
            results = cursor.fetchall()

            if results:
                doc = Document()
                doc.add_heading('Финансовый отчет', level=1)
                doc.add_paragraph(f'Период: {data["start_period"]} - {data["end_period"]}')

                table = doc.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Компания'
                hdr_cells[1].text = 'Сумма'

                for company, total in results:
                    row_cells = table.add_row().cells
                    row_cells[0].text = company
                    row_cells[1].text = str(total)

                file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить отчет", "", "Word Documents (*.docx)")
                if file_path:
                    doc.save(file_path)
                    return file_path
            else:
                QMessageBox.warning(self, 'Ошибка', 'Не удалось найти данные для создания финансового отчета.')

        except pymysql.Error as e:
            QMessageBox.critical(self, 'Ошибка базы данных', f'Ошибка при создании финансового отчета: {e}')
        finally:
            cursor.close()
            connection.close()
        return None

    def logout(self):
        self.close()
        # Логика для возврата к окну входа

class FinancialReportDialog(QDialog):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Сформировать финансовый отчет')
        self.setGeometry(100, 100, 300, 200)

        layout = QFormLayout()

        self.start_period_input = QDateTimeEdit(self)
        self.start_period_input.setCalendarPopup(True)
        self.end_period_input = QDateTimeEdit(self)
        self.end_period_input.setCalendarPopup(True)

        self.company_combobox = QComboBox()

        layout.addRow('Начало периода:', self.start_period_input)
        layout.addRow('Конец периода:', self.end_period_input)
        layout.addRow('Компания:', self.company_combobox)

        self.load_companies()

        self.buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        self.buttons.accepted.connect(self.accept)
        self.buttons.rejected.connect(self.reject)
        layout.addWidget(self.buttons)

        self.setLayout(layout)

    def load_companies(self):
        try:
            connection = connect_to_database()
            cursor = connection.cursor()
            query = "SELECT id, name FROM insurance_company"
            cursor.execute(query)
            companies = cursor.fetchall()
            for company in companies:
                self.company_combobox.addItem(company[1], company[0])
        except pymysql.Error as e:
            QMessageBox.critical(self, 'Ошибка базы данных', f'Ошибка при загрузке компаний: {e}')
        finally:
            cursor.close()
            connection.close()

    def get_data(self):
        return {
            'start_period': self.start_period_input.dateTime().toString('yyyy-MM-dd'),
            'end_period': self.end_period_input.dateTime().toString('yyyy-MM-dd'),
            'company_id': self.company_combobox.currentData(),
        }

def connect_to_database():
    return pymysql.connect(
        host='MySQL-8.2',
        user='root',
        password='',
        database='25.02_ilatovskaya'
    )

if __name__ == '__main__':
    app = QApplication(sys.argv)
    # Пример данных пользователя: (id, name, surname, patronymic, login, password, role_id)
    user_data = (1, 'Иван', 'Иванов', 'Иванович', 'ivanov', 'password123', 3)
    accountant_window = AccountantWindow(user_data=user_data)
    accountant_window.show()
    sys.exit(app.exec())
