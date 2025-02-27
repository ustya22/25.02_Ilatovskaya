import sys
from PyQt6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QLabel, QLineEdit, QPushButton,
    QMessageBox, QCheckBox, QFormLayout, QDialog, QDialogButtonBox, QComboBox,
    QDateTimeEdit, QFileDialog
)
from PyQt6.QtCore import QTimer, Qt
import pymysql
from docx import Document
from docx.shared import Pt

class LabTechnicianWindow(QWidget):
    def __init__(self, user_data):
        super().__init__()
        self.user_data = user_data
        self.session_duration = 10 * 60 * 1000  # 10 минут в миллисекундах
        self.warning_time = 5 * 60 * 1000  # 5 минут в миллисекундах
        self.block_time = 1 * 60 * 1000  # 1 минута в миллисекундах
        self.warning_shown = False  # Флаг для отслеживания показа предупреждения
        self.initUI()
        self.start_session_timer()

    def initUI(self):
        self.setWindowTitle('Лаборант')
        self.setGeometry(100, 100, 400, 300)

        layout = QVBoxLayout()

        self.user_label = QLabel(f'Добро пожаловать, {self.user_data[2]} {self.user_data[1]}')
        layout.addWidget(self.user_label)

        self.role_label = QLabel(f'Роль: {self.get_user_role(self.user_data[6])}')
        layout.addWidget(self.role_label)

        self.timer_label = QLabel("Время сеанса: 00:00")
        layout.addWidget(self.timer_label)

        self.accept_service_button = QPushButton('Зарегистрировать услугу')
        self.accept_service_button.clicked.connect(self.accept_service)
        layout.addWidget(self.accept_service_button)

        self.generate_report_button = QPushButton('Сформировать отчет')
        self.generate_report_button.clicked.connect(self.generate_report)
        layout.addWidget(self.generate_report_button)

        self.logout_button = QPushButton('Выйти')
        self.logout_button.clicked.connect(self.logout)
        layout.addWidget(self.logout_button)

        self.setLayout(layout)

    def get_user_role(self, role_id):
        try:
            connection = connect_to_database()
            cursor = connection.cursor()
            query = "SELECT name FROM post WHERE id = %s"
            cursor.execute(query, (role_id,))
            result = cursor.fetchone()
            return result[0] if result else 'Unknown'
        except pymysql.Error as e:
            QMessageBox.critical(self, 'Database Error', f'Error fetching role: {e}')
            return 'Unknown'
        finally:
            cursor.close()
            connection.close()

    def start_session_timer(self):
        self.time_elapsed = 0
        self.session_timer = QTimer()
        self.session_timer.timeout.connect(self.update_timer)
        self.session_timer.start(1000)  # Обновление каждую секунду
        print("Таймер запущен")  # Отладочное сообщение

    def update_timer(self):
        self.time_elapsed += 1000
        minutes, seconds = divmod(self.time_elapsed // 1000, 60)
        self.timer_label.setText(f"Время сеанса: {minutes:02}:{seconds:02}")
        print(f"Время сеанса: {minutes:02}:{seconds:02}")  # Отладочное сообщение

        if not self.warning_shown and self.time_elapsed >= (self.session_duration - self.warning_time):
            self.show_warning_message()
            self.warning_shown = True  # Установить флаг в True после показа предупреждения

        if self.time_elapsed >= self.session_duration:
            self.session_timer.stop()
            self.end_session()

    def show_warning_message(self):
        QMessageBox.warning(self, 'Предупреждение', 'Сеанс скоро закончится. Пожалуйста, завершите работу.')

    def end_session(self):
        QMessageBox.information(self, 'Сеанс завершен', 'Ваш сеанс закончился. Пожалуйста, выйдите.')
        self.logout()
        self.block_login()

    def block_login(self):
        QTimer.singleShot(self.block_time, self.enable_login)

    def enable_login(self):
        # Логика для разблокировки входа после времени блокировки
        pass

    def logout(self):
        self.close()
        # Логика для возврата к окну входа

    def accept_service(self):
        dialog = ServiceProvidedDialog()
        if dialog.exec():
            data = dialog.get_data()
            self.save_service_to_db(data)

    def save_service_to_db(self, data):
        try:
            connection = connect_to_database()
            cursor = connection.cursor()
            query = """
            INSERT INTO services_provided (id_order, id_employee_services, id_analyzer, start_time, end_time, result)
            VALUES (%s, %s, %s, %s, %s, %s)
            """
            cursor.execute(query, (
                data['id_order'],
                data['id_employee_services'],
                data['id_analyzer'],
                data['start_time'],
                data['end_time'],
                data['result']
            ))
            connection.commit()
            QMessageBox.information(self, 'Успех', 'Услуга зарегистрирована и сохранена в базу данных.')
        except pymysql.Error as e:
            QMessageBox.critical(self, 'Ошибка базы данных', f'Ошибка при сохранении услуги: {e}')
        finally:
            cursor.close()
            connection.close()

    def generate_report(self):
        dialog = ReportDialog()
        if dialog.exec():
            data = dialog.get_data()
            file_path = self.create_report(data)
            if file_path:
                QMessageBox.information(self, 'Отчет', f'Отчет успешно создан и сохранен по пути: {file_path}')

    def create_report(self, data):
        try:
            connection = connect_to_database()
            cursor = connection.cursor()
            query = """
            SELECT p.surname, p.name, p.patronymic, a.name, s.name, sp.start_time, sp.end_time, sp.result
            FROM services_provided sp
            JOIN orders o ON sp.id_order = o.id
            JOIN patient p ON o.id_patient = p.id
            JOIN analyzer a ON sp.id_analyzer = a.id
            JOIN services s ON sp.id_employee_services = s.id
            WHERE sp.id_order = %s AND sp.id_analyzer = %s
            """
            cursor.execute(query, (data['id_order'], data['id_analyzer']))
            result = cursor.fetchone()

            if result:
                doc = Document()
                doc.add_heading('Отчет об анализе', level=1)
                doc.add_paragraph(f'Пациент: {result[0]} {result[1]} {result[2]}')
                doc.add_paragraph(f'Анализатор: {result[3]}')
                doc.add_paragraph(f'Услуга: {result[4]}')
                doc.add_paragraph(f'Время начала: {result[5]}')
                doc.add_paragraph(f'Время окончания: {result[6]}')
                doc.add_paragraph(f'Результат: {result[7]}')

                file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить отчет", "", "Word Documents (*.docx)")
                if file_path:
                    doc.save(file_path)
                    return file_path
            else:
                QMessageBox.warning(self, 'Ошибка', 'Не удалось найти данные для создания отчета.')

        except pymysql.Error as e:
            QMessageBox.critical(self, 'Ошибка базы данных', f'Ошибка при создании отчета: {e}')
        finally:
            cursor.close()
            connection.close()
        return None

class ServiceProvidedDialog(QDialog):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Зарегистрировать услугу')
        self.setGeometry(100, 100, 300, 300)

        layout = QFormLayout()

        self.id_order_combobox = QComboBox()
        self.id_employee_services_combobox = QComboBox()
        self.id_analyzer_combobox = QComboBox()
        self.start_time_input = QDateTimeEdit(self)
        self.start_time_input.setCalendarPopup(True)
        self.end_time_input = QDateTimeEdit(self)
        self.end_time_input.setCalendarPopup(True)
        self.result_input = QLineEdit()

        layout.addRow('Заказ:', self.id_order_combobox)
        layout.addRow('Услуга сотрудника:', self.id_employee_services_combobox)
        layout.addRow('Анализатор:', self.id_analyzer_combobox)
        layout.addRow('Время начала:', self.start_time_input)
        layout.addRow('Время окончания:', self.end_time_input)
        layout.addRow('Результат обследования:', self.result_input)

        self.load_orders()
        self.load_employee_services()
        self.load_analyzers()

        self.buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        self.buttons.accepted.connect(self.accept)
        self.buttons.rejected.connect(self.reject)
        layout.addWidget(self.buttons)

        self.setLayout(layout)

    def load_orders(self):
        try:
            connection = connect_to_database()
            cursor = connection.cursor()
            query = """
            SELECT o.id, CONCAT('Заказ от ', DATE_FORMAT(o.date_of_creation, '%d-%m-%Y'), ' пациента ', p.surname, ' ', p.name)
            FROM orders o
            JOIN patient p ON o.id_patient = p.id
            """
            cursor.execute(query)
            orders = cursor.fetchall()
            for order in orders:
                self.id_order_combobox.addItem(order[1], order[0])
        except pymysql.Error as e:
            QMessageBox.critical(self, 'Ошибка базы данных', f'Ошибка при загрузке заказов: {e}')
        finally:
            cursor.close()
            connection.close()

    def load_employee_services(self):
        try:
            connection = connect_to_database()
            cursor = connection.cursor()
            query = "SELECT id, name FROM services"
            cursor.execute(query)
            services = cursor.fetchall()
            for service in services:
                self.id_employee_services_combobox.addItem(service[1], service[0])
        except pymysql.Error as e:
            QMessageBox.critical(self, 'Ошибка базы данных', f'Ошибка при загрузке услуг: {e}')
        finally:
            cursor.close()
            connection.close()

    def load_analyzers(self):
        try:
            connection = connect_to_database()
            cursor = connection.cursor()
            query = "SELECT id, name FROM analyzer"
            cursor.execute(query)
            analyzers = cursor.fetchall()
            for analyzer in analyzers:
                self.id_analyzer_combobox.addItem(analyzer[1], analyzer[0])
        except pymysql.Error as e:
            QMessageBox.critical(self, 'Ошибка базы данных', f'Ошибка при загрузке анализаторов: {e}')
        finally:
            cursor.close()
            connection.close()

    def get_data(self):
        return {
            'id_order': self.id_order_combobox.currentData(),
            'id_employee_services': self.id_employee_services_combobox.currentData(),
            'id_analyzer': self.id_analyzer_combobox.currentData(),
            'start_time': self.start_time_input.dateTime().toString('yyyy-MM-dd HH:mm:ss'),
            'end_time': self.end_time_input.dateTime().toString('yyyy-MM-dd HH:mm:ss'),
            'result': self.result_input.text(),
        }

class ReportDialog(QDialog):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Сформировать отчет')
        self.setGeometry(100, 100, 300, 200)

        layout = QFormLayout()

        self.id_order_combobox = QComboBox()
        self.id_analyzer_combobox = QComboBox()

        layout.addRow('Заказ:', self.id_order_combobox)
        layout.addRow('Анализатор:', self.id_analyzer_combobox)

        self.load_orders()
        self.load_analyzers()

        self.buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        self.buttons.accepted.connect(self.accept)
        self.buttons.rejected.connect(self.reject)
        layout.addWidget(self.buttons)

        self.setLayout(layout)

    def load_orders(self):
        try:
            connection = connect_to_database()
            cursor = connection.cursor()
            query = """
            SELECT o.id, CONCAT('Заказ от ', DATE_FORMAT(o.date_of_creation, '%d-%m-%Y'), ' пациента ', p.surname, ' ', p.name)
            FROM orders o
            JOIN patient p ON o.id_patient = p.id
            """
            cursor.execute(query)
            orders = cursor.fetchall()
            for order in orders:
                self.id_order_combobox.addItem(order[1], order[0])
        except pymysql.Error as e:
            QMessageBox.critical(self, 'Ошибка базы данных', f'Ошибка при загрузке заказов: {e}')
        finally:
            cursor.close()
            connection.close()

    def load_analyzers(self):
        try:
            connection = connect_to_database()
            cursor = connection.cursor()
            query = "SELECT id, name FROM analyzer"
            cursor.execute(query)
            analyzers = cursor.fetchall()
            for analyzer in analyzers:
                self.id_analyzer_combobox.addItem(analyzer[1], analyzer[0])
        except pymysql.Error as e:
            QMessageBox.critical(self, 'Ошибка базы данных', f'Ошибка при загрузке анализаторов: {e}')
        finally:
            cursor.close()
            connection.close()

    def get_data(self):
        return {
            'id_order': self.id_order_combobox.currentData(),
            'id_analyzer': self.id_analyzer_combobox.currentData(),
        }

def connect_to_database():
    return pymysql.connect(
        host='MySQL-8.2',
        user='root',
        password='',
        database='25.02_ilatovskaya'
    )

if __name__ == '__main__':
    app = QApplication(sys.argv)
    # Пример данных пользователя: (id, name, surname, patronymic, login, password, role_id)
    user_data = (1, 'Иван', 'Иванов', 'Иванович', 'ivanov', 'password123', 1)
    login_window = LabTechnicianWindow(user_data=user_data)
    login_window.show()
    sys.exit(app.exec())
